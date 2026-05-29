"""Markdown -> .docx conversion.

Uses mistune for parsing markdown into an AST and python-docx for emitting
the .docx OOXML. Covers the structures that LLM-generated prompt outputs
actually use: headings, paragraphs, bullet + numbered lists, tables,
fenced code blocks, bold, italic, inline code, and links.

The fidelity bar is "the operator can open it in Word and the structure
is preserved." Not "round-trips through Pandoc."
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import mistune
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


def md_to_docx(md_text: str, output_path: Path) -> None:
    """Convert `md_text` into a .docx file at `output_path`."""
    doc = Document()
    _ensure_code_style(doc)
    parser = mistune.create_markdown(renderer="ast", plugins=["table", "strikethrough"])
    ast = parser(md_text)
    for node in ast:
        _emit_block(doc, node)
    doc.save(str(output_path))


def _ensure_code_style(doc: Document) -> None:
    """Add a 'Code' character style for inline + block code, since python-docx's
    default templates don't ship one in a stable way."""
    styles = doc.styles
    if "InlineCode" not in [s.name for s in styles]:
        s = styles.add_style("InlineCode", WD_STYLE_TYPE.CHARACTER)
        s.font.name = "Courier New"
        s.font.size = Pt(10)


def _emit_block(doc: Document, node: dict[str, Any]) -> None:
    t = node.get("type")
    if t == "heading":
        level = max(1, min(int(node.get("attrs", {}).get("level", 1)), 6))
        para = doc.add_heading(level=level)
        _emit_inlines(para, node.get("children", []))
    elif t == "paragraph":
        para = doc.add_paragraph()
        _emit_inlines(para, node.get("children", []))
    elif t == "list":
        ordered = bool(node.get("attrs", {}).get("ordered"))
        style = "List Number" if ordered else "List Bullet"
        for item in node.get("children", []):
            _emit_list_item(doc, item, style)
    elif t == "block_code":
        text = node.get("raw", "")
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    elif t == "block_quote":
        for child in node.get("children", []):
            _emit_block(doc, child)
    elif t == "table":
        _emit_table(doc, node)
    elif t == "thematic_break":
        doc.add_paragraph("———")
    elif t == "blank_line":
        pass
    else:
        # Unknown block type — render as plain paragraph using flattened text.
        text = _flatten_text(node)
        if text:
            doc.add_paragraph(text)


def _emit_list_item(doc: Document, item: dict[str, Any], style: str) -> None:
    """A list_item contains one or more blocks; render the first paragraph as
    the list bullet, then any further blocks (nested lists, paragraphs) as
    follow-on lines."""
    children = item.get("children", [])
    if not children:
        doc.add_paragraph("", style=style)
        return
    first = children[0]
    para = doc.add_paragraph(style=style)
    if first.get("type") == "block_text" or first.get("type") == "paragraph":
        _emit_inlines(para, first.get("children", []))
        rest = children[1:]
    else:
        _emit_inlines(para, [first])
        rest = children[1:]
    for child in rest:
        _emit_block(doc, child)


def _emit_table(doc: Document, node: dict[str, Any]) -> None:
    rows = node.get("children", [])
    if not rows:
        return
    # First row is the header (mistune's table plugin emits table_head + table_body).
    head_cells: list[list[dict[str, Any]]] = []
    body_rows: list[list[list[dict[str, Any]]]] = []
    for row in rows:
        rtype = row.get("type")
        if rtype == "table_head":
            for cell in row.get("children", []):
                head_cells.append(cell.get("children", []))
        elif rtype == "table_body":
            for body_row in row.get("children", []):
                cells = [c.get("children", []) for c in body_row.get("children", [])]
                body_rows.append(cells)
    if not head_cells:
        return
    table = doc.add_table(rows=1 + len(body_rows), cols=len(head_cells))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for i, hc in enumerate(head_cells):
        para = header_cells[i].paragraphs[0]
        _emit_inlines(para, hc)
        for run in para.runs:
            run.bold = True
    for ri, row_cells in enumerate(body_rows, start=1):
        for ci, cell_inlines in enumerate(row_cells):
            if ci >= len(head_cells):
                break  # malformed row — ignore extras
            para = table.rows[ri].cells[ci].paragraphs[0]
            _emit_inlines(para, cell_inlines)


def _emit_inlines(paragraph: Any, nodes: list[dict[str, Any]]) -> None:
    for node in nodes:
        t = node.get("type")
        if t == "text":
            paragraph.add_run(node.get("raw", ""))
        elif t == "strong":
            run = paragraph.add_run(_flatten_text(node))
            run.bold = True
        elif t == "emphasis":
            run = paragraph.add_run(_flatten_text(node))
            run.italic = True
        elif t == "codespan":
            run = paragraph.add_run(node.get("raw", ""))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif t == "link":
            text = _flatten_text(node)
            url = node.get("attrs", {}).get("url", "")
            run = paragraph.add_run(f"{text} ({url})" if url else text)
            run.font.underline = True
        elif t == "linebreak" or t == "softbreak":
            paragraph.add_run("\n")
        elif t == "strikethrough":
            run = paragraph.add_run(_flatten_text(node))
            run.font.strike = True
        else:
            # Unknown inline — fall back to flattened text.
            text = _flatten_text(node)
            if text:
                paragraph.add_run(text)


def _flatten_text(node: dict[str, Any]) -> str:
    """Recursively concatenate the .raw text of `node` and its descendants."""
    if "raw" in node and isinstance(node["raw"], str):
        return node["raw"]
    parts: list[str] = []
    for child in node.get("children", []) or []:
        parts.append(_flatten_text(child))
    return "".join(parts)
